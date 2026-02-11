#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模版替换引擎 (Python-docx)

使用方式:
    python3 template_replace.py <input.docx> <output.docx> <mappings.json>

mappings.json 格式:
    [
      {"originalText": "原文", "placeholderKey": "field_name"},
      ...
    ]
"""
import sys
import json
import re
import copy
from docx import Document
from docx.oxml.ns import qn

# =========================================================================
# 标签识别
# =========================================================================
KNOWN_LABELS = [
    'SHIPPER', 'CONSIGNEE', 'NOTIFY PARTY',
    'PRE-CARRIAGE BY', 'PLACE OF RECEIPT', 'OCEAN VESSEL/VOY',
    'PORT OF LOADING', 'PORTC OF DISCHARGE', 'PORT OF DISCHARGE',
    'PLACE OF DELIVERY', 'B/L NO.', 'B/L NO',
    'DOC. NO.', 'DOC NO.',
    'SERVICE TYPE / MODE', 'SERVICE TYPE/MODE',
    'LADEN ON BOARD', 'NUMBER OF ORIGINAL B/L(S)',
    'PAYABLE AT', 'PLACE AND DATE OF ISSUE',
    'FREIGHT & CHARGES', 'REVENUE TONS', 'RATE',
    'CONTAINER,SEAL, MARKS & NUMBER', 'QUANTITY AND KIND OF PACKAGES',
    'DESCRIPTION OF GOODS', 'GROSS WEIGHT (KGS)', 'MEASUREMEN(M)',
    'ALSO NOTIFY PARTY (COMPLETE NAME AND ADDRESS)',
    'BOOKING NO.',
]

# 这些文本既可以是列标题（标签），也可以是值
# 仅当它们出现在单元格第一段时才被视为标签
AMBIGUOUS_LABELS = ['PREPAID', 'COLLECT', 'FREIGHT PREPAID', 'AS ARRANGED', 'TELEX RELEASE']

LABEL_TO_KEY = {
    'SHIPPER': 'shipper',
    'CONSIGNEE': 'consignee',
    'NOTIFY PARTY': 'notify_party',
    'PRE-CARRIAGE BY': 'pre_carriage_by',
    'PLACE OF RECEIPT': 'place_of_receipt',
    'OCEAN VESSEL/VOY': 'vessel_voyage',
    'PORT OF LOADING': 'port_of_loading',
    'PORTC OF DISCHARGE': 'port_of_discharge',
    'PORT OF DISCHARGE': 'port_of_discharge',
    'PLACE OF DELIVERY': 'place_of_delivery',
    'B/L NO.': 'bl_no',
    'B/L NO': 'bl_no',
    'DOC. NO.': 'doc_no',
    'DOC NO.': 'doc_no',
    'BOOKING NO.': 'booking_no',
    'SERVICE TYPE / MODE': 'service_type',
    'SERVICE TYPE/MODE': 'service_type',
    'PAYABLE AT': 'payable_at',
    'PLACE AND DATE OF ISSUE': 'issue_place',
    'CONTAINER,SEAL, MARKS & NUMBER': 'container_no',
    'REVENUE TONS': 'revenue_tons',
    'NUMBER OF ORIGINAL B/L(S)': 'original_bl_count',
    'FREIGHT & CHARGES': 'freight_term',
    'DESCRIPTION OF GOODS': 'goods_description',
    'GROSS WEIGHT (KGS)': 'gross_weight_kgs',
    'MEASUREMEN(M)': 'measurement_cbm',
    'QUANTITY AND KIND OF PACKAGES': 'package_quantity',
}


def is_label(text, is_first_paragraph=True):
    """判断文本是否为结构性标签（不应被替换的标题/列头）
    
    is_first_paragraph: 是否是单元格的第一个段落。
    某些文本（如 PREPAID, AS ARRANGED）仅在作为第一段时才是标签。
    """
    if not text or not text.strip():
        return False
    upper = text.strip().upper()
    for label in KNOWN_LABELS:
        # 严格匹配：只有完全等于标签，或者 "标签:" 才是标签
        # "标签 值" 不应被视为标签，而应视为包含标签的值行
        if upper == label or upper == label + ':' or upper == label + ' :':
            return True
    # 歧义标签：仅在第一段时视为标签
    if is_first_paragraph:
        for label in AMBIGUOUS_LABELS:
            if upper == label:
                return True
    # 模式匹配：以 "FOR DELIVERY" / "ONWARD INLAND" / "In Witness" 开头的固定文本
    # 如果包含 ":" 且长度超过40，可能是 "Label: Value" 的混合行，不应视为纯标签（应视为值以便替换）
    is_mixed_line = ':' in text and len(text) > 40

    if (upper.startswith('FOR DELIVERY') or
        upper.startswith('ONWARD INLAND') or
        upper.startswith('ALSO NOTIFY')) and not is_mixed_line:
        return True
    if text.strip().startswith('In Witness') or text.strip().startswith('ofLading'):
        return True
    return False


def get_label_key(label):
    """根据标签文本获取对应的 placeholder_key"""
    if not label:
        return None
    upper = label.strip().upper().replace(':', '').strip()
    # 精确匹配
    if upper in LABEL_TO_KEY:
        return LABEL_TO_KEY[upper]
    # 前缀匹配 (归一化去除空格后再比对)
    norm = re.sub(r'[^A-Z0-9]+', '', upper)
    if norm.startswith('FORDELIVERY') or norm.startswith('ALSONOTIFY'):
        return 'delivery_agent'
    return None


def normalize(text):
    """归一化文本：换行→空格，多余空白→单个空格"""
    if not text: return ""
    return re.sub(r'\s+', ' ', text.replace('\n', ' ')).strip()


def normalize_compare(text):
    """极致归一化：仅保留字母数字，用于抗干扰对比"""
    if not text: return ""
    return re.sub(r'[^a-zA-Z0-9]+', '', text).lower()


# =========================================================================
# 段落级别的精准替换（保留 Run 样式）
# =========================================================================
def surgical_replace_in_paragraph(paragraph, target, replacement):
    """在段落中精确替换 target → replacement，保留尽可能多的样式"""
    runs = paragraph.runs
    if not runs:
        return False

    # 1. 构建全文本和 Run 边界坐标
    full_text = ''
    run_boundaries = []  # [(start, end, run_index)]
    for i, run in enumerate(runs):
        t = run.text or ''
        start = len(full_text)
        full_text += t
        run_boundaries.append((start, len(full_text), i))

    # 2. 查找目标
    match_idx = full_text.find(target)
    if match_idx == -1:
        return False

    match_end = match_idx + len(target)

    # 3. 找到受影响的 Run 范围
    start_run_idx = None
    end_run_idx = None
    for s, e, idx in run_boundaries:
        if match_idx >= s and match_idx < e:
            start_run_idx = idx
        if match_end > s and match_end <= e:
            end_run_idx = idx

    if start_run_idx is None or end_run_idx is None:
        return False

    # 4. 执行替换
    start_run = runs[start_run_idx]
    start_text = start_run.text or ''
    offset_in_start = match_idx - run_boundaries[start_run_idx][0]

    if start_run_idx == end_run_idx:
        # 目标在单个 Run 内
        prefix = start_text[:offset_in_start]
        suffix = start_text[offset_in_start + len(target):]
        start_run.text = prefix + replacement + suffix
    else:
        # 跨多个 Run
        prefix = start_text[:offset_in_start]
        start_run.text = prefix + replacement

        # 清空末尾 Run 的受影响部分
        end_run = runs[end_run_idx]
        end_text = end_run.text or ''
        offset_in_end = match_end - run_boundaries[end_run_idx][0]
        end_run.text = end_text[offset_in_end:]

        # 删除中间的 Run（从后向前）
        for i in range(end_run_idx - 1, start_run_idx, -1):
            paragraph._p.remove(runs[i]._r)

    return True


def clear_and_set_paragraph(paragraph, new_text):
    """清空段落所有 Run 并设置新文本（保留第一个 Run 的样式）"""
    runs = paragraph.runs
    if not runs:
        if new_text:
            run = paragraph.add_run(new_text)
        return

    # 保留第一个 Run 的样式
    runs[0].text = new_text
    # 删除其余 Run（从后向前）
    for i in range(len(runs) - 1, 0, -1):
        paragraph._p.remove(runs[i]._r)


# =========================================================================
# 特殊处理：包含 delivery_agent + carrier_agent 的复合单元格
# =========================================================================
def process_complex_delivery_cell(cell, mappings_by_key, sorted_mappings):
    """
    处理包含多个逻辑区域的复合单元格（如提单中的交付/签署区）。
    结构通常如下：
      P[0]:  "FOR DELIVERY OF GOODS PLEASE APPLY TO: <delivery_agent_line1>"
      P[1-N]: <delivery_agent 后续行>
      P[...]: ""  (空行)
      P[...]: "ALSO NOTIFY PARTY (COMPLETE NAME AND ADDRESS)"
      P[...]: ""  (空行)
      P[...]: "TELEX RELEASE"
      P[...]: ""  (空行)
      P[...]: "In Witness Whereof, ... on behalf of <carrier_agent> Has signed..."
      P[...]: "ofLading stated below..."

    如果不是这种模式，返回 False，由标准 process_cell 处理。
    """
    paragraphs = cell.paragraphs
    if not paragraphs or len(paragraphs) < 5:
        return False

    # 检测模式：第一个非空段落以 "FOR DELIVERY" 开头
    first_text = ''
    for p in paragraphs:
        if p.text and p.text.strip():
            first_text = p.text.strip()
            break
    if not first_text.upper().startswith('FOR DELIVERY'):
        return False

    # ===== 处理 delivery_agent =====
    if 'delivery_agent' in mappings_by_key:
        placeholder = '{{delivery_agent}}'
        # 找到 delivery_agent 值的范围：从 P[0] 的冒号后面开始，到第一个空段落或 "ALSO NOTIFY" 为止
        delivery_start_idx = 0
        delivery_end_idx = 0  # exclusive

        for i, p in enumerate(paragraphs):
            text = p.text.strip().upper() if p.text else ''
            if i > 0 and (not text or text.startswith('ALSO NOTIFY')):
                delivery_end_idx = i
                break
        else:
            delivery_end_idx = len(paragraphs)

        # P[0]: 保留 "FOR DELIVERY OF GOODS PLEASE APPLY TO: " 前缀，替换其后内容
        if delivery_end_idx > 0:
            p0 = paragraphs[delivery_start_idx]
            p0_text = p0.text.strip()
            if ':' in p0_text:
                prefix = p0_text.split(':', 1)[0] + ': '
                clear_and_set_paragraph(p0, prefix + placeholder)
            else:
                clear_and_set_paragraph(p0, placeholder)

            # 清空 P[1] 到 P[delivery_end_idx-1]（delivery agent 的后续行）
            for j in range(delivery_start_idx + 1, delivery_end_idx):
                clear_and_set_paragraph(paragraphs[j], '')

    # ===== 处理 carrier_agent =====
    if 'carrier_agent' in mappings_by_key:
        placeholder = '{{carrier_agent}}'
        # carrier_agent 嵌在 "In Witness Whereof, ... on behalf of <NAME> Has signed..." 中
        for i, p in enumerate(paragraphs):
            text = p.text.strip() if p.text else ''
            if text.startswith('In Witness'):
                # 提取 "on behalf of <NAME> Has signed" 中的 <NAME>
                upper = text.upper()
                behalf_idx = upper.find('ON BEHALF OF')
                has_idx = upper.find('HAS SIGNED')
                if behalf_idx != -1 and has_idx != -1 and has_idx > behalf_idx:
                    prefix_part = text[:behalf_idx + len('ON BEHALF OF')]
                    suffix_part = text[has_idx:]
                    new_text = prefix_part + ' ' + placeholder + ' ' + suffix_part
                    clear_and_set_paragraph(p, new_text)
                break

    return True


# =========================================================================
# 核心：以单元格为单位处理
# =========================================================================
def process_cell(cell, mappings_by_key, sorted_mappings):
    """
    处理单个单元格：
    1. 先检查是否为复合单元格（delivery_agent + carrier_agent），若是则使用专用处理器
    2. 识别标签段落 vs 值段落
    3. 标签直接查找占位符（消歧）
    4. 回退到内容精确匹配
    5. 替换值段落
    """
    paragraphs = cell.paragraphs
    if not paragraphs:
        return

    # 优先检测复合单元格
    if process_complex_delivery_cell(cell, mappings_by_key, sorted_mappings):
        # 仍然递归处理嵌套表格
        for nested_table in cell.tables:
            for row in nested_table.rows:
                for nested_cell in row.cells:
                    process_cell(nested_cell, mappings_by_key, sorted_mappings)
        return

    # --- 第一步：分离标签和值 ---
    first_value_idx = -1
    cell_label = ''
    found_first_non_empty = False

    for i, p in enumerate(paragraphs):
        text = p.text
        if not text or not text.strip():
            continue
        trimmed = text.strip()
        is_first = not found_first_non_empty
        found_first_non_empty = True
        
        if is_label(trimmed, is_first_paragraph=is_first) and first_value_idx == -1:
            cell_label = trimmed
            continue
        if first_value_idx == -1:
            first_value_idx = i

    if first_value_idx == -1:
        # 特殊情况：有标签但没有值（值段落为空）
        # 尝试寻找标签后的第一个段落（即使是空的）作为值段落
        if cell_label:
             for i, p in enumerate(paragraphs):
                 text = p.text
                 if not text or not text.strip():
                     # 找到一个空段落，且在标签之后（通常标签是第一段）
                     # 这里简单假设标签后的第一个空段落就是值
                     if i > 0: 
                         first_value_idx = i
                         break
        
        if first_value_idx == -1:
            return  # 确实无值段落

    # --- 第二步：拼接所有值段落 ---
    value_parts = []
    value_indexes = []
    for i in range(first_value_idx, len(paragraphs)):
        text = paragraphs[i].text
        if not text or not text.strip():
            continue
        # 值段落中不跳过歧义标签（如 FREIGHT PREPAID, AS ARRANGED）
        if is_label(text.strip(), is_first_paragraph=False):
            continue
        value_parts.append(text.strip())
        value_indexes.append(i)

    cell_value = ' '.join(value_parts)
    if not cell_value:
        return

    normalized_cell = normalize(cell_value)

    # --- 第三步：匹配 ---
    matched_placeholder = None

    # 策略A：标签直接查找
    if cell_label:
        label_key = get_label_key(cell_label)
        if label_key and label_key in mappings_by_key:
            matched_placeholder = '{{' + label_key + '}}'

    # 策略B：内容精确匹配（回退）
    if not matched_placeholder:
        for m in sorted_mappings:
            orig = m['originalText']
            normalized_orig = normalize_compare(orig)
            
            # 使用极致归一化对比
            cell_compare = normalize_compare(normalized_cell)

            if cell_compare == normalized_orig:
                matched_placeholder = '{{' + m['placeholderKey'] + '}}'
                break
            # 允许短数字匹配 (如 "68")，非数字需 >= 3 chars
            elif normalized_orig in cell_compare and (len(normalized_orig) >= 3 or normalized_orig.isdigit()):
                # 特殊逻辑：如果差异仅在于前缀（FOR DELIVERY...），则视为全匹配，触发后续的前缀保留逻辑
                if cell_compare.endswith(normalized_orig):
                    # 检查 prefix 关键词 (norm_cell_prefix)
                    prefix_norm = cell_compare[:-len(normalized_orig)]
                    if "fordelivery" in prefix_norm or "alsonotify" in prefix_norm:
                         matched_placeholder = '{{' + m['placeholderKey'] + '}}'
                         break

                # 子串匹配 — 手术刀替换
                placeholder = '{{' + m['placeholderKey'] + '}}'
                # 寻找原文中第1行作为锚点 (避开换行符干扰)
                first_line = orig.split('\n')[0].strip()
                for idx in value_indexes:
                    p = paragraphs[idx]
                    # 如果精确包含则直接换；如果不包含，尝试在这个段落里模糊搜一下
                    if p.text and first_line in p.text:
                        surgical_replace_in_paragraph(p, first_line, placeholder)
                        break
                    elif p.text and normalize_compare(first_line) in normalize_compare(p.text):
                         # 这里的模糊匹配较危险，仅在长名称文本时尝试
                         if len(first_line) > 10:
                             surgical_replace_in_paragraph(p, p.text, placeholder) # 整个段落换掉
                             break
                continue  # 不消耗整个单元格

    # 策略C：逐段落精确匹配（处理短文本如 N/M）
    # 当整个单元格无法整体匹配时，尝试逐段落与映射值精确比对
    if not matched_placeholder:
        for m in sorted_mappings:
            orig = m['originalText']
            if not orig:
                continue
            first_line = orig.split('\n')[0].strip()
            if not first_line:
                continue
            placeholder = '{{' + m['placeholderKey'] + '}}'
            for idx in value_indexes:
                p = paragraphs[idx]
                p_text = p.text.strip() if p.text else ''
                # 段落文本与映射原文精确匹配
                if p_text and p_text == first_line:
                    surgical_replace_in_paragraph(p, first_line, placeholder)
                    break

    # --- 第四步：整体替换 ---
    if matched_placeholder:
        is_first = True
        for idx in value_indexes:
            if is_first:
                p = paragraphs[idx]
                p_text = p.text.strip()
                p_upper = p_text.upper()
                new_text = matched_placeholder

                # 特殊处理：保留"混合行"的前缀（如 "FOR DELIVERY...: "）
                # 只有当该行被判定为值（非标签）且拥有特定前缀时触发
                check_prefix = p_upper.startswith('FOR DELIVERY') or p_upper.startswith('ALSO NOTIFY')
                if check_prefix:
                    if ':' in p_text:
                        # 保留冒号前的内容
                        parts = p_text.split(':', 1)
                        new_text = parts[0] + ': ' + matched_placeholder
                    elif ' TO ' in p_upper or p_upper.endswith(' TO'):
                        # 保留 TO 前的内容
                        idx_to = p_upper.rfind('TO')
                        new_text = p_text[:idx_to+2] + ' ' + matched_placeholder

                clear_and_set_paragraph(p, new_text)
                is_first = False
            else:
                clear_and_set_paragraph(paragraphs[idx], '')

    # 递归处理嵌套表格
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                process_cell(nested_cell, mappings_by_key, sorted_mappings)


def replace_in_paragraph(paragraph, sorted_mappings):
    """在正文段落（非单元格）中做简单替换"""
    text = paragraph.text
    if not text or not text.strip():
        return
    for m in sorted_mappings:
        target = m['originalText']
        if not target:
            continue
        first_line = target.split('\n')[0].strip()
        if first_line in text:
            surgical_replace_in_paragraph(paragraph, first_line, '{{' + m['placeholderKey'] + '}}')
            text = paragraph.text  # 更新


# =========================================================================
# 主入口
# =========================================================================
def process_document(input_path, output_path, mappings):
    """
    主处理函数：
    input_path:  原始 .docx 文件路径
    output_path: 输出 .docx 文件路径
    mappings:    映射列表 [{"originalText": "...", "placeholderKey": "..."}, ...]
    """
    doc = Document(input_path)

    if mappings is None:
        mappings = []

    # 过滤无效映射
    valid = [m for m in mappings
             if m.get('originalText') and m.get('placeholderKey')]

    # 按原文长度倒序
    valid.sort(key=lambda m: -len(m['originalText']))

    # key → mapping 索引
    by_key = {}
    for m in valid:
        by_key[m['placeholderKey']] = m

    # 1. 处理所有表格（以单元格为单位）
    # 用一个字典保存 id(tc) -> tc，防止对象被 GC 导致 ID 复用
    tc_vault = {} 
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                tc = cell._tc
                cid = id(tc)
                if cid in tc_vault:
                    continue
                tc_vault[cid] = tc
                
                process_cell(cell, by_key, valid)

    # 2. 处理正文段落
    for p in doc.paragraphs:
        replace_in_paragraph(p, valid)

    # 3. 处理页眉/页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for p in header.paragraphs:
                    replace_in_paragraph(p, valid)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for p in footer.paragraphs:
                    replace_in_paragraph(p, valid)

    doc.save(output_path)
    print(json.dumps({"status": "success", "output": output_path}))


if __name__ == '__main__':
    if len(sys.argv) != 4:
        print(json.dumps({"status": "error", "message": "Usage: template_replace.py <input> <output> <mappings.json>"}))
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    mappings_file = sys.argv[3]

    try:
        with open(mappings_file, 'r', encoding='utf-8') as f:
            mappings_data = json.load(f)

        # 兼容 Java 传入的驼峰命名 (originalText / placeholderKey)
        # 以及 Dify 返回的蛇形命名 (original_text / placeholder_key)
        normalized_mappings = []
        for m in mappings_data:
            nm = {
                'originalText': m.get('originalText') or m.get('original_text') or '',
                'placeholderKey': m.get('placeholderKey') or m.get('placeholder_key') or '',
            }
            normalized_mappings.append(nm)

        process_document(input_file, output_file, normalized_mappings)
    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}))
        sys.exit(1)
